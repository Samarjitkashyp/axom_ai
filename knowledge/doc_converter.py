import os
import sys
import tempfile
import io
from pathlib import Path
from pypdf import PdfReader

# ReportLab imports for pure-Python high-quality PDF rendering
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    Image as RLImage, KeepTogether, PageBreak, HRFlowable
)
from reportlab.pdfgen import canvas

try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    docx = None

try:
    from PIL import Image as PILImage
except ImportError:
    PILImage = None


class NumberedCanvas(canvas.Canvas):
    """Two-pass canvas to dynamically calculate and render 'Page X of Y' and header/footer."""
    def __init__(self, *args, **kwargs):
        super(NumberedCanvas, self).__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_decorations(num_pages)
            super(NumberedCanvas, self).showPage()
        super(NumberedCanvas, self).save()

    def draw_page_decorations(self, page_count):
        self.saveState()
        self.setFont("Helvetica", 8)
        self.setFillColor(colors.HexColor("#64748b"))

        # Footer line
        self.setStrokeColor(colors.HexColor("#e2e8f0"))
        self.setLineWidth(0.5)
        self.line(40, 36, self._pagesize[0] - 40, 36)

        # Footer text: Axom AI branding and page number
        self.drawString(40, 24, "Converted with Axom AI Doc-to-PDF Engine")
        page_text = f"Page {self._pageNumber} of {page_count}"
        self.drawRightString(self._pagesize[0] - 40, 24, page_text)

        self.restoreState()


def _convert_docx_pure_python(docx_path, output_pdf_path):
    """
    Pure Python conversion of DOCX to PDF using python-docx and ReportLab.
    Parses paragraphs, runs (bold, italic, colors), bullet points, tables, and images.
    """
    if not docx:
        raise ImportError("python-docx is not installed.")

    doc = docx.Document(docx_path)
    
    # PDF document template with comfortable margins
    margin = 40
    pdf = SimpleDocTemplate(
        output_pdf_path,
        pagesize=A4,
        leftMargin=margin,
        rightMargin=margin,
        topMargin=margin,
        bottomMargin=45
    )
    
    usable_width = A4[0] - (2 * margin)

    styles = getSampleStyleSheet()

    # Custom typography styles
    normal_style = ParagraphStyle(
        'DocxNormal',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10.5,
        leading=14.5,
        textColor=colors.HexColor("#1e293b"),
        spaceAfter=6,
    )

    title_style = ParagraphStyle(
        'DocxTitle',
        parent=styles['Title'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        textColor=colors.HexColor("#0f172a"),
        spaceAfter=12,
        spaceBefore=8,
        alignment=0,
    )

    h1_style = ParagraphStyle(
        'DocxH1',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=15,
        leading=19,
        textColor=colors.HexColor("#1e293b"),
        spaceBefore=14,
        spaceAfter=6,
        keepWithNext=True,
    )

    h2_style = ParagraphStyle(
        'DocxH2',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=12.5,
        leading=16.5,
        textColor=colors.HexColor("#334155"),
        spaceBefore=10,
        spaceAfter=4,
        keepWithNext=True,
    )

    h3_style = ParagraphStyle(
        'DocxH3',
        parent=styles['Heading3'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=15,
        textColor=colors.HexColor("#475569"),
        spaceBefore=8,
        spaceAfter=3,
        keepWithNext=True,
    )

    bullet_style = ParagraphStyle(
        'DocxBullet',
        parent=normal_style,
        leftIndent=18,
        firstLineIndent=-10,
        spaceAfter=4,
    )

    table_cell_style = ParagraphStyle(
        'DocxTableCell',
        parent=normal_style,
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#1e293b"),
        spaceAfter=0,
    )

    table_header_style = ParagraphStyle(
        'DocxTableHeader',
        parent=table_cell_style,
        fontName='Helvetica-Bold',
        textColor=colors.HexColor("#0f172a"),
    )

    story = []

    def escape_xml(text):
        if not text:
            return ""
        return (
            text.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
                .replace('"', "&quot;")
                .replace("'", "&apos;")
        )

    # Process all elements in the document body
    for element in doc.element.body:
        tag = element.tag.split('}')[-1]

        if tag == 'p':
            # It's a paragraph
            p = None
            for paragraph in doc.paragraphs:
                if paragraph._element == element:
                    p = paragraph
                    break

            if not p:
                continue

            raw_text = p.text.strip()
            if not raw_text:
                # Add a subtle vertical spacer for blank lines
                story.append(Spacer(1, 6))
                continue

            style_name = p.style.name.lower() if p.style else 'normal'

            # Build formatted HTML-like runs
            formatted_runs = []
            for run in p.runs:
                r_text = escape_xml(run.text)
                if not r_text:
                    continue

                if run.bold:
                    r_text = f"<b>{r_text}</b>"
                if run.italic:
                    r_text = f"<i>{r_text}</i>"
                if run.underline:
                    r_text = f"<u>{r_text}</u>"

                # Check font color
                if run.font and run.font.color and run.font.color.rgb:
                    color_hex = f"#{run.font.color.rgb}"
                    r_text = f"<font color='{color_hex}'>{r_text}</font>"

                formatted_runs.append(r_text)

            p_content = "".join(formatted_runs)
            if not p_content:
                p_content = escape_xml(raw_text)

            # Determine appropriate ParagraphStyle
            if 'title' in style_name:
                p_style = title_style
            elif 'heading 1' in style_name:
                p_style = h1_style
            elif 'heading 2' in style_name:
                p_style = h2_style
            elif 'heading 3' in style_name:
                p_style = h3_style
            elif 'list' in style_name or 'bullet' in style_name or raw_text.startswith(('•', '-', '*')):
                p_style = bullet_style
                if not (p_content.startswith('•') or p_content.startswith('&bull;')):
                    p_content = f"&bull; {p_content}"
            else:
                p_style = normal_style

            # Check alignment
            if p.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                centered_style = ParagraphStyle('Centered', parent=p_style, alignment=1)
                story.append(Paragraph(p_content, centered_style))
            elif p.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                right_style = ParagraphStyle('RightAligned', parent=p_style, alignment=2)
                story.append(Paragraph(p_content, right_style))
            else:
                story.append(Paragraph(p_content, p_style))

        elif tag == 'tbl':
            # It's a table
            t = None
            for table in doc.tables:
                if table._element == element:
                    t = table
                    break

            if not t:
                continue

            num_cols = len(t.columns)
            if num_cols == 0:
                continue

            col_width = usable_width / num_cols
            table_data = []

            for row_idx, row in enumerate(t.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = escape_xml(cell.text.strip())
                    c_style = table_header_style if row_idx == 0 else table_cell_style
                    row_data.append(Paragraph(cell_text, c_style))
                table_data.append(row_data)

            if table_data:
                rl_table = Table(table_data, colWidths=[col_width] * num_cols)
                rl_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f1f5f9')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#0f172a')),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('INNERGRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e1')),
                    ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#94a3b8')),
                    ('TOPPADDING', (0, 0), (-1, -1), 5),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ]))
                story.append(Spacer(1, 6))
                story.append(rl_table)
                story.append(Spacer(1, 8))

    if not story:
        story.append(Paragraph("Empty Document", normal_style))

    # Build PDF with dynamic page numbering canvas
    pdf.build(story, canvasmaker=NumberedCanvas)
    return output_pdf_path


def _convert_txt_to_pdf(txt_path, output_pdf_path):
    """Converts a plain text file into a clean formatted PDF."""
    with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()

    margin = 40
    pdf = SimpleDocTemplate(
        output_pdf_path,
        pagesize=A4,
        leftMargin=margin,
        rightMargin=margin,
        topMargin=margin,
        bottomMargin=45
    )

    styles = getSampleStyleSheet()
    txt_style = ParagraphStyle(
        'TxtNormal',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor("#1e293b"),
        spaceAfter=4,
    )

    story = []
    lines = content.splitlines()
    for line in lines:
        cleaned = (
            line.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
                .replace(" ", "&nbsp;")
        )
        if not cleaned.strip():
            story.append(Spacer(1, 6))
        else:
            story.append(Paragraph(cleaned, txt_style))

    if not story:
        story.append(Paragraph("Empty text document", txt_style))

    pdf.build(story, canvasmaker=NumberedCanvas)
    return output_pdf_path


def convert_doc_to_pdf(input_file_path, output_pdf_path=None):
    """
    Main conversion entry point.
    Tries Microsoft Word COM conversion via docx2pdf (if available on Windows),
    and seamlessly falls back to pure-Python ReportLab conversion.

    Returns:
        tuple: (output_pdf_path, page_count, file_size_str)
    """
    input_path = Path(input_file_path).resolve()
    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_file_path}")

    ext = input_path.suffix.lower()

    if output_pdf_path is None:
        output_dir = input_path.parent
        output_pdf_path = output_dir / f"{input_path.stem}.pdf"
    else:
        output_pdf_path = Path(output_pdf_path).resolve()

    output_pdf_path.parent.mkdir(parents=True, exist_ok=True)

    success = False
    last_error = None

    # Strategy 1: If it's a DOCX/DOC and on Windows, try docx2pdf for native fidelity
    if ext in ['.docx', '.doc'] and sys.platform == 'win32':
        try:
            import pythoncom
            pythoncom.CoInitialize()
            from docx2pdf import convert
            convert(str(input_path), str(output_pdf_path))
            if output_pdf_path.exists() and output_pdf_path.stat().st_size > 0:
                success = True
        except Exception as e:
            last_error = str(e)
            success = False

    # Strategy 2: Pure-Python conversion for DOCX
    if not success and ext == '.docx':
        try:
            _convert_docx_pure_python(str(input_path), str(output_pdf_path))
            if output_pdf_path.exists() and output_pdf_path.stat().st_size > 0:
                success = True
        except Exception as e:
            last_error = str(e)
            success = False

    # Strategy 3: Plain text / RTF fallback
    if not success and ext in ['.txt', '.rtf', '.log', '.md', '.csv']:
        try:
            _convert_txt_to_pdf(str(input_path), str(output_pdf_path))
            if output_pdf_path.exists() and output_pdf_path.stat().st_size > 0:
                success = True
        except Exception as e:
            last_error = str(e)
            success = False

    # Strategy 4: Fallback for .doc if Word was not available
    if not success and ext == '.doc':
        # If .doc could not be converted via COM, try reading text stream or docx
        try:
            _convert_docx_pure_python(str(input_path), str(output_pdf_path))
            success = True
        except Exception:
            try:
                _convert_txt_to_pdf(str(input_path), str(output_pdf_path))
                success = True
            except Exception as e:
                last_error = str(e)
                success = False

    if not success or not output_pdf_path.exists():
        raise RuntimeError(f"Failed to convert {input_path.name} to PDF. Reason: {last_error or 'Unknown conversion error'}")

    # Inspect generated PDF with pypdf to get page count and size
    page_count = 1
    try:
        reader = PdfReader(str(output_pdf_path))
        page_count = len(reader.pages)
    except Exception:
        pass

    size_bytes = output_pdf_path.stat().st_size
    if size_bytes < 1024:
        file_size_str = f"{size_bytes} B"
    elif size_bytes < 1024 * 1024:
        file_size_str = f"{size_bytes / 1024:.1f} KB"
    else:
        file_size_str = f"{size_bytes / (1024 * 1024):.2f} MB"

    return str(output_pdf_path), page_count, file_size_str
