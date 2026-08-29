import os
import csv
import re
import json
import difflib
import requests
import numpy as np
from django.db.models import Q
from .models import KnowledgeDocument, KnowledgeChunk, QAPair

# --------------------------------------------------------------------------
# Semantic search — BAAI/bge-m3 (multilingual, runs locally, no rate limits).
# The model is loaded lazily once and kept resident in the process.
# --------------------------------------------------------------------------
EMBED_MODEL = os.getenv('EMBED_MODEL', 'BAAI/bge-m3')
SEMANTIC_THRESHOLD = float(os.getenv('SEMANTIC_THRESHOLD', '0.72'))
_EMBED_MODEL_OBJ = None
_QA_CACHE = None  # (ids, answers, normalized_matrix, count)


def _get_model():
    """Lazy-load the sentence-transformers model once (kept in memory)."""
    global _EMBED_MODEL_OBJ
    if _EMBED_MODEL_OBJ is None:
        from sentence_transformers import SentenceTransformer
        _EMBED_MODEL_OBJ = SentenceTransformer(EMBED_MODEL)
    return _EMBED_MODEL_OBJ


def _embed_texts(texts):
    """Embed a list of texts locally with bge-m3. Returns a list of normalized
    vectors, or None on failure (caller falls back to keyword search)."""
    if not texts:
        return None
    try:
        model = _get_model()
        vecs = model.encode(texts, normalize_embeddings=True, batch_size=32)
        return [v.tolist() for v in vecs]
    except Exception:
        return None


def backfill_qa_embeddings(batch_size=256):
    """Compute + store embeddings for every QAPair that doesn't have one yet."""
    global _QA_CACHE
    pending = list(QAPair.objects.filter(embedding='').only('id', 'question'))
    done = 0
    for i in range(0, len(pending), batch_size):
        batch = pending[i:i + batch_size]
        vecs = _embed_texts([qa.question for qa in batch])
        if not vecs or len(vecs) != len(batch):
            continue
        for qa, v in zip(batch, vecs):
            qa.embedding = json.dumps(v)
        QAPair.objects.bulk_update(batch, ['embedding'])
        done += len(batch)
    _QA_CACHE = None
    return done


def _load_qa_matrix():
    """Load all QAPair embeddings into a normalized numpy matrix (cached)."""
    global _QA_CACHE
    total = QAPair.objects.exclude(embedding='').count()
    if _QA_CACHE is not None and _QA_CACHE[5] == total:
        return _QA_CACHE
    ids, answers, asms, srcs, vecs = [], [], [], [], []
    for qa in QAPair.objects.exclude(embedding='').only(
            'id', 'answer', 'answer_assamese', 'embedding', 'source_name', 'source_url').iterator():
        try:
            vecs.append(json.loads(qa.embedding))
            ids.append(qa.id)
            answers.append(qa.answer)
            asms.append(qa.answer_assamese)
            srcs.append({'name': qa.source_name, 'url': qa.source_url})
        except Exception:
            continue
    if not vecs:
        _QA_CACHE = ([], [], [], [], None, 0)
        return _QA_CACHE
    mat = np.asarray(vecs, dtype=np.float32)
    norms = np.linalg.norm(mat, axis=1, keepdims=True)
    norms[norms == 0] = 1.0
    _QA_CACHE = (ids, answers, asms, srcs, mat / norms, total)
    return _QA_CACHE


def semantic_find_answer(query, threshold=None):
    """Return (answer, answer_assamese, score, source) for the semantically closest
    stored question, or (None, '', score, None) if nothing is confident enough.
    `source` is {'name':..., 'url':...} (empty strings when the pair has no source).
    The answer is verbatim from the knowledge base, so it stays factually consistent."""
    if threshold is None:
        threshold = SEMANTIC_THRESHOLD
    ids, answers, asms, srcs, mat, total = _load_qa_matrix()
    if mat is None:
        return None, '', 0.0, None
    qv = _embed_texts([query])
    if not qv:
        return None, '', 0.0, None
    q = np.asarray(qv[0], dtype=np.float32)
    n = np.linalg.norm(q)
    if n == 0:
        return None, '', 0.0, None
    sims = mat @ (q / n)
    idx = int(np.argmax(sims))
    score = float(sims[idx])
    if score >= threshold:
        return answers[idx], asms[idx], score, srcs[idx]
    return None, '', score, None

def extract_text_from_file(file_path, file_type):
    """
    Extracts plain text and structured content from PDF, Excel, Image, or Text files.
    """
    text_content = ""
    file_type = file_type.lower()

    try:
        # PDF Extraction using pypdf
        if file_type == 'pdf' or file_path.endswith('.pdf'):
            import pypdf
            reader = pypdf.PdfReader(file_path)
            extracted_pages = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    extracted_pages.append(f"--- Page {i+1} ---\n{page_text}")
            text_content = "\n\n".join(extracted_pages)

        # JSONL / JSON Extraction — training-style Q&A pairs made searchable
        elif file_type == 'jsonl' or file_path.endswith(('.jsonl', '.json')):
            import json
            blocks = []

            def format_obj(obj):
                if isinstance(obj, dict):
                    if 'instruction' in obj and 'output' in obj:
                        return f"Q: {obj['instruction']}\nA: {obj['output']}"
                    if 'question' in obj and 'answer' in obj:
                        return f"Q: {obj['question']}\nA: {obj['answer']}"
                    return "\n".join(f"{k}: {v}" for k, v in obj.items())
                return str(obj)

            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                raw = f.read().strip()

            # Try whole-file JSON first (a JSON array), else parse line-by-line (JSONL).
            parsed_whole = None
            try:
                parsed_whole = json.loads(raw)
            except Exception:
                parsed_whole = None

            if isinstance(parsed_whole, list):
                for obj in parsed_whole:
                    blocks.append(format_obj(obj))
            elif isinstance(parsed_whole, dict):
                blocks.append(format_obj(parsed_whole))
            else:
                for line in raw.splitlines():
                    line = line.strip()
                    if not line:
                        continue
                    try:
                        blocks.append(format_obj(json.loads(line)))
                    except Exception:
                        blocks.append(line)

            text_content = "\n\n".join(blocks)

        # Word DOCX Extraction using python-docx (paragraphs + tables)
        elif file_type == 'docx' or file_path.endswith('.docx'):
            import docx
            document = docx.Document(file_path)
            blocks = [p.text for p in document.paragraphs if p.text and p.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                    if cells:
                        blocks.append(" | ".join(cells))
            text_content = "\n".join(blocks)

        # Excel / CSV Extraction using openpyxl or csv
        elif file_type in ['excel', 'csv', 'xlsx', 'xls'] or file_path.endswith(('.xlsx', '.xls', '.csv')):
            if file_path.endswith('.csv'):
                with open(file_path, mode='r', encoding='utf-8', errors='ignore') as f:
                    reader = csv.reader(f)
                    rows = [", ".join(row) for row in reader if row]
                    text_content = "\n".join(rows)
            else:
                import openpyxl
                wb = openpyxl.load_workbook(file_path, data_only=True)
                sheet_data = []
                for sheet in wb.sheetnames:
                    ws = wb[sheet]
                    sheet_data.append(f"=== Sheet: {sheet} ===")
                    for row in ws.iter_rows(values_only=True):
                        row_vals = [str(val) for val in row if val is not None]
                        if row_vals:
                            sheet_data.append(" | ".join(row_vals))
                text_content = "\n".join(sheet_data)

        # Image Metadata & Description using Pillow
        elif file_type in ['image', 'png', 'jpg', 'jpeg', 'webp'] or file_path.endswith(('.png', '.jpg', '.jpeg', '.webp')):
            from PIL import Image
            with Image.open(file_path) as img:
                filename = os.path.basename(file_path)
                text_content = f"Image File: {filename}\nFormat: {img.format}\nDimensions: {img.width}x{img.height} px\nColor Mode: {img.mode}"

        # Plain Text / Markdown / JSON Files
        else:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text_content = f.read()

    except Exception as e:
        text_content = f"Error extracting content: {str(e)}"

    return text_content.strip()


def create_knowledge_chunks(document_obj, chunk_size=800, overlap=100):
    """
    Splits document extracted text into chunks and stores KnowledgeChunk DB objects.
    """
    text = document_obj.extracted_text
    if not text:
        return

    document_obj.chunks.all().delete()

    chunks_data = []
    start = 0
    chunk_index = 0

    while start < len(text):
        end = start + chunk_size
        chunk_text = text[start:end]
        
        chunk_obj = KnowledgeChunk(
            document=document_obj,
            content=chunk_text,
            chunk_index=chunk_index,
            keywords=" ".join(set(chunk_text.lower().split()[:20]))
        )
        chunks_data.append(chunk_obj)
        
        start += (chunk_size - overlap)
        chunk_index += 1

    KnowledgeChunk.objects.bulk_create(chunks_data)


def create_qa_pairs(document_obj):
    """
    Build QAPair rows from an uploaded document, then embed them so they are
    immediately searchable.

    JSONL/JSON files may give each line the richest control — every object can
    carry: question, answer, answer_assamese (optional), source_name/source_url
    (optional; falls back to the document's own source). Other file types fall
    back to parsing 'Q: ...\\nA: ...' blocks from the extracted text and inherit
    the document-level source.
    """
    document_obj.qa_pairs.all().delete()
    doc_src_name = (document_obj.source_name or '').strip()
    doc_src_url = (document_obj.source_url or '').strip()

    pairs = []

    # 1. Try structured JSONL/JSON first — one JSON object per line (or a JSON array).
    is_jsonish = (document_obj.file_type == 'jsonl'
                  or str(document_obj.file.name).lower().endswith(('.jsonl', '.json')))
    if is_jsonish:
        try:
            with open(document_obj.file.path, 'r', encoding='utf-8') as f:
                raw = f.read()
            records = []
            stripped = raw.strip()
            if stripped.startswith('['):          # a JSON array
                records = json.loads(stripped)
            else:                                  # JSONL — one object per line
                for line in stripped.splitlines():
                    line = line.strip()
                    if line:
                        records.append(json.loads(line))
            for obj in records:
                if not isinstance(obj, dict):
                    continue
                question = str(obj.get('question') or obj.get('instruction') or '').strip()
                answer = str(obj.get('answer') or obj.get('output') or '').strip()
                if not (question and answer):
                    continue
                pairs.append(QAPair(
                    document=document_obj,
                    question=question,
                    answer=answer,
                    answer_assamese=str(obj.get('answer_assamese') or '').strip(),
                    source_name=str(obj.get('source_name') or doc_src_name).strip()[:255],
                    source_url=str(obj.get('source_url') or doc_src_url).strip()[:500],
                ))
        except Exception:
            pairs = []  # fall through to the plain-text parser below

    # 2. Fallback: 'Q: ...\nA: ...' blocks from the extracted text.
    if not pairs:
        text = document_obj.extracted_text or ""
        for block in text.split("\n\n"):
            block = block.strip()
            if block.startswith("Q:") and "\nA:" in block:
                q_part, a_part = block.split("\nA:", 1)
                question = q_part[2:].strip()
                answer = a_part.strip()
                if question and answer:
                    pairs.append(QAPair(
                        document=document_obj, question=question, answer=answer,
                        source_name=doc_src_name[:255], source_url=doc_src_url[:500],
                    ))

    if not pairs:
        return

    QAPair.objects.bulk_create(pairs)

    # 3. Embed the new pairs right away so semantic search can find them.
    try:
        fresh = list(document_obj.qa_pairs.filter(embedding='').only('id', 'question'))
        vecs = _embed_texts([qa.question for qa in fresh])
        if vecs and len(vecs) == len(fresh):
            for qa, v in zip(fresh, vecs):
                qa.embedding = json.dumps(v)
            QAPair.objects.bulk_update(fresh, ['embedding'])
    except Exception:
        pass  # embeddings can be backfilled later via manage.py backfill_embeddings

    global _QA_CACHE
    _QA_CACHE = None  # force the search matrix to rebuild with the new pairs


def _normalize(s):
    """Lowercase, strip punctuation and collapse whitespace for matching."""
    return re.sub(r'\s+', ' ', re.sub(r'[^\w\s]', '', s.lower())).strip()


# Common grammatical / question words that carry little topical meaning. Keeping
# these out of matching lets different phrasings of the same question line up.
_STOPWORDS = {
    'kya', 'hai', 'hain', 'tha', 'the', 'ka', 'ke', 'ki', 'ko', 'me', 'mein', 'se',
    'aur', 'ek', 'ye', 'yeh', 'wo', 'woh', 'iska', 'uska', 'par', 'hi', 'bhi',
    'kaise', 'kaun', 'kaunsa', 'kaunsi', 'kab', 'kahan', 'kyun', 'kitna',
    'hota', 'hoti', 'hote', 'kar', 'karo', 'karta', 'batao', 'bata', 'baare', 'samjhao',
    'is', 'are', 'was', 'a', 'an', 'of', 'to', 'in', 'on', 'the', 'for', 'and', 'about',
    'what', 'which', 'who', 'how', 'when', 'where', 'why', 'tell', 'does', 'do', 'me', 'my',
    # conversational filler — so greetings/chit-chat don't accidentally match KB text
    'you', 'your', 'yours', 'am', 'be', 'been', 'being', 'doing', 'going', 'get',
    'hello', 'hii', 'hey', 'namaste', 'thanks', 'thank', 'please', 'okay', 'yes', 'no',
    'ho', 'hu', 'hoon', 'aap', 'tum', 'raha', 'rahe', 'rahi', 'good', 'nice', 'help',
}


def _keywords(s):
    """Meaningful lowercase tokens (drops stopwords and very short words)."""
    return {t for t in re.sub(r'[^\w\s]', ' ', s.lower()).split()
            if len(t) > 2 and t not in _STOPWORDS}


def find_instant_answer(query, threshold=0.6):
    """
    Return a stored answer when the query shares enough meaningful keywords with
    a saved question. This is phrasing-tolerant ("Bihu kab hota hai?", "When is
    Bihu?", "Rongali Bihu kya hai?" all match) and fast even with many pairs:
    the database pre-filters to questions sharing at least one keyword, then we
    score only those. Returns None when nothing is confident enough, so the
    caller falls back to RAG + the language model.
    """
    q_tokens = _keywords(query)
    if not q_tokens:
        return None, '', None

    qn = _normalize(query)

    # DB-side pre-filter: only pull questions that share a keyword (fast at scale).
    q_filter = Q()
    for t in q_tokens:
        q_filter |= Q(question__icontains=t)
    candidates = QAPair.objects.filter(q_filter).only(
        'question', 'answer', 'answer_assamese', 'source_name', 'source_url')[:3000]

    best_answer, best_asm, best_src = None, '', None
    best_score = 0.0
    for qa in candidates:
        if _normalize(qa.question) == qn:
            # exact match — instant
            return qa.answer, qa.answer_assamese, {'name': qa.source_name, 'url': qa.source_url}

        p_tokens = _keywords(qa.question)
        if not p_tokens:
            continue
        inter = len(q_tokens & p_tokens)
        if inter == 0:
            continue
        # Overlap coefficient (tolerant to extra words), tie-broken by Jaccard
        # so a tighter, more specific question wins when coverage is equal.
        overlap = inter / min(len(q_tokens), len(p_tokens))
        jaccard = inter / len(q_tokens | p_tokens)
        score = overlap + jaccard * 0.001
        if score > best_score:
            best_score = score
            best_answer = qa.answer
            best_asm = qa.answer_assamese
            best_src = {'name': qa.source_name, 'url': qa.source_url}

    if best_score >= threshold:
        return best_answer, best_asm, best_src
    return None, '', None


def search_knowledge_base(query, top_k=3):
    """
    High-performance database query checking if user question exists in Knowledge Base.
    Returns (context_string, source_doc_titles).
    """
    if not query:
        return "", []

    clean_query = query.strip()
    # Only meaningful keywords (stopwords like "kaise", "kya" dropped) so an
    # off-topic query with common filler words doesn't falsely match context.
    tokens = list(_keywords(clean_query))
    if not tokens:
        return "", []

    # Fast OR Query filtering on Database Chunks
    q_objects = Q()
    for token in tokens:
        q_objects |= Q(content__icontains=token) | Q(keywords__icontains=token)

    chunks = KnowledgeChunk.objects.filter(q_objects).select_related('document')[:top_k]

    if not chunks.exists():
        # Fallback: match a document by TITLE only (precise) — not full text, so a
        # single common word can't drag in an unrelated document's content.
        for token in tokens:
            docs = KnowledgeDocument.objects.filter(title__icontains=token)[:1]
            if docs.exists():
                d = docs[0]
                return f"[Document: {d.title}]\n{d.extracted_text[:1000]}", [d.title]
        return "", []

    context_blocks = []
    titles = set()
    for chunk in chunks:
        context_blocks.append(f"[Document: {chunk.document.title}]\n{chunk.content}")
        titles.add(chunk.document.title)

    return "\n\n".join(context_blocks), list(titles)
